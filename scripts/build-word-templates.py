from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "templates"

WORD_FILENAMES = {
    "software-hardware-interface-list": "软硬件接口清单",
    "report-template-list": "报告模板清单",
    "pre-go-live-task-plan": "上线前任务计划",
    "environment-deployment-record": "环境部署记录",
    "initial-configuration-record": "初始化配置记录",
    "interface-document-confirmation": "接口文档确认记录",
    "sequence-diagram-review": "时序图（评审确认版）",
    "interface-integration-record": "接口联调记录",
    "interface-acceptance-form": "接口对接确认表",
    "end-to-end-test-record": "全流程内测记录",
    "training-record": "培训记录",
    "trial-test-record": "试行测试记录",
    "go-live-plan": "上线方案",
    "go-live-confirmation": "上线确认单",
}

# compact_reference_guide preset with named project overrides:
# A4 landscape, 0.60/0.65-inch margins, Microsoft YaHei, and blue-white
# healthcare implementation colors.
FONT = "Microsoft YaHei"
PRIMARY = "1769E0"
NAVY = "17365D"
BODY = "24364B"
MUTED = "667A90"
PALE_BLUE = "EDF5FF"
PALE_BLUE_2 = "F6F9FD"
PALE_YELLOW = "FFF9E8"
LINE = "C8D8EA"
WHITE = "FFFFFF"
SUCCESS = "E4F4EC"
TEAL = "0F766E"
PURPLE = "6D4AA2"
ORANGE = "C66A16"
RED = "B33A3A"
GREEN = "2F7D52"
GOLD = "B88917"
PALE_TEAL = "EAF6F4"
PALE_PURPLE = "F3EFF9"
PALE_ORANGE = "FFF3E8"
PALE_RED = "FCEEEE"
PALE_GREEN = "EDF7F1"

PAGE_WIDTH_IN = 11.69
PAGE_HEIGHT_IN = 8.27
MARGIN_TOP_IN = 0.60
MARGIN_BOTTOM_IN = 0.60
MARGIN_SIDE_IN = 0.65
CONTENT_WIDTH_DXA = 14900
TABLE_INDENT_DXA = 120


SPECS = [
    {
        "slug": "software-hardware-interface-list",
        "stage": "事前准备",
        "layout": "inventory",
        "accent": TEAL,
        "title": "软硬件接口清单",
        "subtitle": "用于项目启动阶段识别系统接口、服务器、终端、网络与外设依赖",
        "instruction": "由项目经理牵头，实施、研发、客户信息部门共同确认。数量、版本、部署位置和责任人应可追溯。",
        "summary_labels": ["清单范围", "盘点日期", "牵头部门", "确认轮次"],
        "columns": ["序号", "接口/设备名称", "类型", "规格/版本", "数量", "部署位置", "责任人", "确认状态", "备注"],
        "weights": [0.55, 1.55, 0.90, 1.10, 0.55, 1.30, 0.85, 0.90, 1.30],
        "rows": 9,
        "checks": [
            ("网络与安全", "端口、访问策略、证书和账号已确认"),
            ("设备与资源", "型号、数量、安装位置和供货责任已确认"),
            ("接口依赖", "协议、联调前置条件和对接联系人已明确"),
        ],
    },
    {
        "slug": "report-template-list",
        "stage": "事前准备",
        "layout": "catalog",
        "accent": PURPLE,
        "title": "报告模板清单",
        "subtitle": "用于确认项目需交付的报告、单据、标签和统计模板",
        "instruction": "逐项确认模板来源、适用业务、样例版本和业务确认人；涉及打印的模板需同步验证纸张与打印机设置。",
        "summary_labels": ["业务范围", "模板总数", "业务确认人", "确认轮次"],
        "columns": ["序号", "模板名称", "适用业务", "版本", "来源", "负责人", "确认状态", "备注"],
        "weights": [0.55, 1.55, 1.35, 0.70, 0.95, 0.90, 0.90, 1.40],
        "rows": 9,
        "checks": [
            ("内容完整性", "字段、口径、排序及页眉页脚已确认"),
            ("打印适配", "纸张、方向、边距、条码及打印机已验证"),
            ("版本管理", "正式版本、样例文件和变更记录可追溯"),
        ],
    },
    {
        "slug": "pre-go-live-task-plan",
        "stage": "事前准备",
        "layout": "countdown",
        "accent": ORANGE,
        "title": "上线前任务计划",
        "subtitle": "用于统筹上线前的准备事项、前置依赖、完成标准与责任人",
        "instruction": "任务应拆分到可验证的完成标准。计划日期变化、阻塞原因和风险升级应在备注中留痕。",
        "summary_labels": ["计划窗口", "任务总数", "总负责人", "评审日期"],
        "columns": ["序号", "任务项", "负责人", "计划开始", "计划结束", "前置条件", "完成标准", "状态", "备注"],
        "weights": [0.50, 1.50, 0.85, 0.90, 0.90, 1.25, 1.40, 0.80, 1.30],
        "rows": 10,
        "checks": [
            ("范围冻结", "上线范围、版本、数据和接口边界已冻结"),
            ("资源到位", "人员、账号、环境、设备和应急联系方式齐备"),
            ("准入评审", "未完成项已评估，不影响上线准入结论"),
        ],
    },
    {
        "slug": "environment-deployment-record",
        "stage": "事前准备",
        "layout": "deployment",
        "accent": TEAL,
        "title": "环境部署记录",
        "subtitle": "用于记录各环境的部署基线、版本、执行人和验证结果",
        "instruction": "每次部署应对应明确版本或制品标识；服务器地址、口令等敏感信息请仅记录保管位置，不在本表明文填写。",
        "summary_labels": ["部署环境", "部署窗口", "实施负责人", "变更单号"],
        "columns": ["序号", "环境名称", "服务器/地址", "部署组件", "部署版本", "部署时间", "实施人", "验证结果", "备注"],
        "weights": [0.50, 0.90, 1.35, 1.20, 1.00, 1.05, 0.80, 0.90, 1.30],
        "rows": 8,
        "checks": [
            ("部署前", "备份、空间、账号权限和依赖项检查完成"),
            ("部署后", "服务、日志、访问、接口和关键业务冒烟验证完成"),
            ("可回退", "备份位置、回退步骤和执行责任人已确认"),
        ],
    },
    {
        "slug": "initial-configuration-record",
        "stage": "事前准备",
        "layout": "configuration",
        "accent": PURPLE,
        "title": "初始化配置记录",
        "subtitle": "用于记录系统上线前的基础参数、业务规则与验证结果",
        "instruction": "配置值应与已确认需求一致；涉及账号、密钥、个人信息等敏感内容仅记录引用编号或保管位置。",
        "summary_labels": ["配置环境", "配置批次", "配置负责人", "复核日期"],
        "columns": ["序号", "配置域", "配置项", "配置内容/引用", "配置人", "配置时间", "验证方式", "验证结果", "备注"],
        "weights": [0.50, 0.95, 1.30, 1.50, 0.80, 1.00, 1.20, 0.85, 1.05],
        "rows": 8,
        "checks": [
            ("需求一致", "配置项与确认需求、数据字典和权限方案一致"),
            ("双人复核", "关键参数已由非配置人员复核并记录结论"),
            ("配置留档", "导出文件、截图或变更记录已归档"),
        ],
    },
    {
        "slug": "interface-document-confirmation",
        "stage": "接口对接",
        "layout": "technical-review",
        "accent": NAVY,
        "title": "接口文档确认记录",
        "subtitle": "用于确认接口协议、字段、编码、异常处理和版本基线",
        "instruction": "确认结论应指向具体接口文档版本；未决字段、异常码和业务口径必须形成责任人与完成日期。",
        "summary_labels": ["接口范围", "文档版本", "提出方", "确认方"],
        "columns": ["序号", "接口名称", "确认内容", "提出方", "确认方", "确认日期", "确认结论", "备注"],
        "weights": [0.50, 1.25, 2.20, 0.85, 0.85, 0.95, 0.95, 1.25],
        "rows": 8,
        "checks": [
            ("协议与安全", "地址、协议、鉴权、加密和重试策略已确认"),
            ("数据口径", "字段、类型、长度、必填、字典和时间格式已确认"),
            ("异常处理", "错误码、超时、补偿、幂等和日志追踪已确认"),
        ],
    },
    {
        "slug": "sequence-diagram-review",
        "stage": "接口对接",
        "layout": "diagram-review",
        "accent": PURPLE,
        "title": "时序图（评审确认版）",
        "subtitle": "用于沉淀跨系统业务时序、关键节点、异常分支与评审结论",
        "instruction": "请将最终时序图粘贴至图示区或作为附件引用，并在评审记录中明确版本、结论和未决事项。",
        "summary_labels": ["业务场景", "图示版本", "参与系统", "评审日期"],
        "columns": ["序号", "评审要点", "涉及系统/角色", "评审意见", "责任人", "计划完成", "状态"],
        "weights": [0.50, 1.60, 1.35, 2.20, 0.85, 0.95, 0.85],
        "rows": 6,
        "checks": [
            ("正常链路", "触发条件、调用顺序、响应和落库节点清晰"),
            ("异常分支", "超时、失败、重试、回退和人工补偿已覆盖"),
            ("一致性", "时序图与接口文档、业务流程和部署方案一致"),
        ],
        "diagram_box": True,
    },
    {
        "slug": "interface-integration-record",
        "stage": "接口对接",
        "layout": "integration",
        "accent": TEAL,
        "title": "接口联调记录",
        "subtitle": "用于记录接口联调批次、测试数据、问题整改与最终确认",
        "instruction": "每条记录应能关联接口版本、测试数据和问题编号；敏感测试数据需脱敏，不在附件中保留真实身份信息。",
        "summary_labels": ["联调环境", "联调批次", "双方负责人", "执行日期"],
        "columns": ["序号", "接口名称", "用例/场景", "测试数据引用", "测试结果", "问题编号", "整改结果", "确认人", "日期"],
        "weights": [0.50, 1.15, 1.45, 1.25, 0.90, 0.90, 1.20, 0.80, 0.85],
        "rows": 9,
        "checks": [
            ("可追溯", "接口版本、请求标识、日志时间和问题编号可关联"),
            ("问题闭环", "失败场景已有根因、责任人、修复版本和复测结论"),
            ("联调完成", "双方已确认正常、异常及边界场景的联调范围"),
        ],
    },
    {
        "slug": "interface-acceptance-form",
        "stage": "接口对接",
        "layout": "acceptance",
        "accent": GREEN,
        "title": "接口对接确认表",
        "subtitle": "用于接口阶段验收，确认范围、标准、遗留问题和责任边界",
        "instruction": "仅在关键用例通过、问题处置明确且双方认可的情况下签署；有条件通过时应列明限制条件和关闭日期。",
        "summary_labels": ["验收范围", "接口版本", "验收环境", "验收日期"],
        "columns": ["序号", "验收项目", "验收标准", "验收结果", "问题/说明", "责任人", "完成日期", "签字/确认"],
        "weights": [0.50, 1.25, 1.80, 0.85, 1.65, 0.85, 0.90, 1.10],
        "rows": 8,
        "checks": [
            ("范围完整", "验收接口、版本、环境和用例范围无遗漏"),
            ("质量达标", "关键用例通过，性能、安全和稳定性满足约定"),
            ("遗留可控", "遗留问题不阻断业务，责任与关闭计划明确"),
        ],
    },
    {
        "slug": "end-to-end-test-record",
        "stage": "上线试运行",
        "layout": "test-charter",
        "accent": PURPLE,
        "title": "全流程内测记录",
        "subtitle": "用于验证端到端业务场景、角色权限、数据流与外围系统协同",
        "instruction": "优先覆盖高频、关键和异常场景；测试数据应脱敏，失败用例需关联问题编号并完成复测。",
        "summary_labels": ["测试版本", "测试环境", "测试负责人", "测试日期"],
        "columns": ["序号", "业务场景", "前置条件/数据", "测试步骤", "预期结果", "实际结果", "结论", "问题编号", "测试人"],
        "weights": [0.45, 1.20, 1.35, 1.70, 1.25, 1.20, 0.70, 0.80, 0.75],
        "rows": 9,
        "checks": [
            ("业务覆盖", "核心、异常、权限和跨系统场景覆盖充分"),
            ("数据安全", "测试数据已脱敏，截图和日志不含敏感信息"),
            ("缺陷闭环", "阻断和严重问题已关闭，其余问题已有处置计划"),
        ],
    },
    {
        "slug": "training-record",
        "stage": "上线试运行",
        "layout": "training",
        "accent": TEAL,
        "title": "培训记录",
        "subtitle": "用于记录培训计划、内容、参训人员、效果评价与后续安排",
        "instruction": "按岗位区分培训内容；签到、课件、现场照片或线上会议记录可作为附件归档。",
        "summary_labels": ["培训主题", "培训对象", "培训方式", "培训日期"],
        "columns": ["序号", "姓名/工号", "部门/岗位", "联系方式", "签到", "考核/反馈", "需跟进事项", "确认签字"],
        "weights": [0.50, 1.10, 1.25, 1.15, 0.75, 1.25, 1.75, 1.10],
        "rows": 10,
        "checks": [
            ("内容适配", "不同岗位的操作、权限和常见问题已覆盖"),
            ("效果确认", "关键岗位已完成实操或理解度确认"),
            ("资料归档", "课件、签到、答疑和补训计划已留档"),
        ],
    },
    {
        "slug": "trial-test-record",
        "stage": "上线试运行",
        "layout": "trial-observation",
        "accent": ORANGE,
        "title": "试行测试记录",
        "subtitle": "用于记录试运行期间的业务使用、问题反馈、处置与关闭情况",
        "instruction": "按科室或岗位记录真实业务试行情况；涉及患者或业务敏感信息时，仅填写脱敏标识或问题编号。",
        "summary_labels": ["试行范围", "试行版本", "试行负责人", "试行周期"],
        "columns": ["序号", "科室/岗位", "试行场景", "试行时间", "试行结果", "问题与建议", "跟进人", "关闭情况", "备注"],
        "weights": [0.45, 1.00, 1.35, 0.95, 0.90, 1.70, 0.80, 0.90, 1.10],
        "rows": 9,
        "checks": [
            ("业务可用", "关键岗位能够独立完成日常业务操作"),
            ("运行稳定", "系统、接口、设备和打印在试行期运行稳定"),
            ("问题可控", "遗留问题有分级、责任人、计划和应急方案"),
        ],
    },
    {
        "slug": "go-live-plan",
        "stage": "上线试运行",
        "layout": "cutover-command",
        "accent": RED,
        "title": "上线方案",
        "subtitle": "用于组织上线窗口、执行步骤、责任分工、风险应对与回退",
        "instruction": "上线步骤应按时间顺序可执行、可验证、可回退；关键联系人在上线窗口内保持可联络。",
        "summary_labels": ["上线范围", "上线版本", "上线窗口", "总指挥"],
        "columns": ["阶段", "上线事项", "执行步骤/检查点", "负责人", "计划时间", "完成标志", "回退方案", "状态"],
        "weights": [0.75, 1.25, 2.20, 0.85, 1.00, 1.25, 1.55, 0.75],
        "rows": 10,
        "checks": [
            ("上线准入", "版本、数据、环境、人员和业务确认均已完成"),
            ("应急保障", "监控、备份、回退触发条件和应急联系人明确"),
            ("上线收口", "验证、通知、值守和上线后复盘安排已明确"),
        ],
    },
    {
        "slug": "go-live-confirmation",
        "stage": "上线试运行",
        "layout": "formal-confirmation",
        "accent": GREEN,
        "title": "上线确认单",
        "subtitle": "用于上线准入或上线完成后的多方确认与责任留痕",
        "instruction": "确认结果分为通过、有条件通过和不通过；有条件通过时必须填写限制条件、责任人和关闭日期。",
        "summary_labels": ["上线项目", "上线版本", "确认类型", "确认日期"],
        "columns": ["序号", "确认项", "确认标准", "确认结果", "说明/证据", "责任方", "完成日期", "签字/确认"],
        "weights": [0.50, 1.25, 1.85, 0.85, 1.65, 0.85, 0.90, 1.10],
        "rows": 8,
        "checks": [
            ("技术确认", "部署、接口、数据、备份、监控和安全验证通过"),
            ("业务确认", "关键流程、权限、模板、设备和用户培训满足上线"),
            ("管理确认", "遗留事项、保障安排和责任边界已经各方确认"),
        ],
    },
]


def set_run_font(run, size: float, color: str = BODY, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, size: float = 9.2, color: str = BODY, bold: bool = False,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def normalize_widths(weights: Sequence[float], total: int = CONTENT_WIDTH_DXA) -> list[int]:
    denominator = sum(weights)
    widths = [round(total * weight / denominator) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: Sequence[int], total: int = CONTENT_WIDTH_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = LINE, size: str = "5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=8.5, color=MUTED)


LAYOUT_LABELS = {
    "inventory": "资源盘点",
    "catalog": "模板目录",
    "countdown": "上线倒排",
    "deployment": "部署作业",
    "configuration": "配置治理",
    "technical-review": "技术评审",
    "diagram-review": "时序评审",
    "integration": "联调执行",
    "acceptance": "接口验收",
    "test-charter": "场景测试",
    "training": "培训组织",
    "trial-observation": "试行观察",
    "cutover-command": "上线指挥",
    "formal-confirmation": "正式确认",
}


def accent_fill(accent: str) -> str:
    return {
        TEAL: PALE_TEAL,
        PURPLE: PALE_PURPLE,
        ORANGE: PALE_ORANGE,
        RED: PALE_RED,
        GREEN: PALE_GREEN,
    }.get(accent, PALE_BLUE)


def configure_document(doc: Document, spec: dict) -> None:
    title = spec["title"]
    accent = spec["accent"]
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_TOP_IN)
    section.bottom_margin = Inches(MARGIN_BOTTOM_IN)
    section.left_margin = Inches(MARGIN_SIDE_IN)
    section.right_margin = Inches(MARGIN_SIDE_IN)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, PRIMARY, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(10.2))
    set_table_geometry(table, normalize_widths([1.0, 1.0]))
    set_table_borders(table, color=WHITE, size="0")
    set_cell_text(table.cell(0, 0), f"实施项目管理平台 · {LAYOUT_LABELS[spec['layout']]}", size=8.5, color=accent, bold=True)
    set_cell_text(table.cell(0, 1), f"{title}  |  V2.1", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_number(p)

    doc.core_properties.title = title
    doc.core_properties.subject = "项目实施 SOP 交付物模板"
    doc.core_properties.author = "实施项目管理平台"
    doc.core_properties.comments = f"V2.1 {LAYOUT_LABELS[spec['layout']]}差异化模板"


def add_title_block(doc: Document, spec: dict) -> None:
    accent = spec["accent"]
    layout = spec["layout"]
    if layout in {"countdown", "cutover-command"}:
        banner = doc.add_table(rows=1, cols=1)
        banner.style = "Table Grid"
        set_table_geometry(banner, [CONTENT_WIDTH_DXA])
        set_table_borders(banner, color=accent, size="0")
        cell = banner.cell(0, 0)
        set_cell_fill(cell, accent)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(f"{spec['title']}\n")
        set_run_font(run, size=23, color=WHITE, bold=True)
        run = paragraph.add_run(spec["subtitle"])
        set_run_font(run, size=10.5, color=WHITE)
    else:
        centered = layout in {"diagram-review", "acceptance", "formal-confirmation"}
        align = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        title = doc.add_paragraph()
        title.alignment = align
        title.paragraph_format.space_before = Pt(3)
        title.paragraph_format.space_after = Pt(3)
        title.paragraph_format.keep_with_next = True
        run = title.add_run(spec["title"])
        set_run_font(run, size=25 if centered else 23, color=NAVY, bold=True)

        subtitle = doc.add_paragraph()
        subtitle.alignment = align
        subtitle.paragraph_format.space_before = Pt(0)
        subtitle.paragraph_format.space_after = Pt(9)
        subtitle.paragraph_format.keep_with_next = True
        run = subtitle.add_run(spec["subtitle"])
        set_run_font(run, size=10.5, color=MUTED)

    meta = doc.add_table(rows=2, cols=8)
    meta.style = "Table Grid"
    widths = normalize_widths([0.72, 1.55, 0.72, 1.55, 0.72, 1.35, 0.72, 1.67])
    set_table_geometry(meta, widths)
    set_table_borders(meta)
    labels = ["项目名称", "客户单位", "项目经理", "文档编号", "实施负责人", "编制日期", "模板版本", "文档状态"]
    values = ["", "", "", "", "", "", "V2.1", f"☐ 草稿  ☐ {LAYOUT_LABELS[layout]}完成"]
    for row_index in range(2):
        for pair_index in range(4):
            logical_index = row_index * 4 + pair_index
            label_cell = meta.cell(row_index, pair_index * 2)
            value_cell = meta.cell(row_index, pair_index * 2 + 1)
            set_cell_fill(label_cell, accent_fill(accent))
            set_cell_fill(value_cell, WHITE if values[logical_index] else PALE_YELLOW)
            set_cell_text(label_cell, labels[logical_index], size=9, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(value_cell, values[logical_index], size=9, color=BODY)


def add_callout(doc: Document, text: str, accent: str = PRIMARY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent)
    cell = table.cell(0, 0)
    set_cell_fill(cell, accent_fill(accent))
    set_cell_text(cell, f"填写提示  |  {text}", size=9.1, color=BODY)
    cell.paragraphs[0].runs[0].bold = False


def add_section_title(doc: Document, title: str, accent: str = PRIMARY,
                      page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(title)
    set_run_font(run, size=13, color=accent, bold=True)


def add_summary_table(doc: Document, labels: Sequence[str]) -> None:
    table = doc.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    set_table_geometry(table, normalize_widths([0.75, 1.55] * 4))
    set_table_borders(table)
    for row_index in range(2):
        for pair_index in range(4):
            logical_index = row_index * 4 + pair_index
            label = labels[logical_index] if logical_index < len(labels) else ["归档位置", "附件数量", "复核人", "复核日期"][logical_index - len(labels)]
            set_cell_fill(table.cell(row_index, pair_index * 2), PALE_BLUE)
            set_cell_fill(table.cell(row_index, pair_index * 2 + 1), PALE_YELLOW)
            set_cell_text(table.cell(row_index, pair_index * 2), label, size=8.9, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(table.cell(row_index, pair_index * 2 + 1), "", size=8.9)


OVERVIEW_METRICS = {
    "inventory": [("软件与接口", "版本 / 协议"), ("服务器终端", "型号 / 数量"), ("网络安全", "地址 / 端口"), ("外设打印", "驱动 / 纸张")],
    "catalog": [("业务报告", "字段与口径"), ("单据表单", "版式与签署"), ("标签条码", "尺寸与打印"), ("统计模板", "周期与权限")],
    "countdown": [("D-14", "范围冻结"), ("D-7", "环境就绪"), ("D-3", "全流程验证"), ("D-1", "上线准入")],
    "deployment": [("01", "备份与空间"), ("02", "制品与依赖"), ("03", "部署与迁移"), ("04", "冒烟与回退")],
    "configuration": [("基础参数", "组织 / 字典"), ("业务规则", "流程 / 编码"), ("权限安全", "角色 / 审计"), ("输出设置", "模板 / 打印")],
    "technical-review": [("协议", "地址 / 鉴权"), ("数据", "字段 / 字典"), ("异常", "重试 / 补偿"), ("版本", "基线 / 变更")],
    "diagram-review": [("触发方", "业务入口"), ("参与系统", "调用边界"), ("关键节点", "落库 / 回执"), ("异常分支", "重试 / 人工")],
    "integration": [("请求", "报文与标识"), ("响应", "状态与耗时"), ("日志", "时间与链路"), ("缺陷", "修复与复测")],
    "acceptance": [("范围", "接口与版本"), ("质量", "功能与稳定"), ("遗留", "限制与计划"), ("结论", "通过条件")],
    "test-charter": [("核心流程", "端到端业务"), ("角色权限", "岗位与数据"), ("外围协同", "接口与设备"), ("异常恢复", "失败与补偿")],
    "training": [("对象", "岗位分组"), ("内容", "操作与规则"), ("实操", "任务与考核"), ("跟进", "答疑与补训")],
    "trial-observation": [("业务可用", "岗位独立操作"), ("运行稳定", "系统接口设备"), ("用户反馈", "问题与建议"), ("上线准备", "遗留与保障")],
    "cutover-command": [("D-1", "最终检查"), ("T0", "变更开始"), ("T+30", "业务验证"), ("T+120", "稳定确认")],
    "formal-confirmation": [("技术确认", "部署与数据"), ("业务确认", "流程与人员"), ("保障确认", "监控与值守"), ("管理结论", "责任与日期")],
}


OVERVIEW_DETAILS = {
    "inventory": [("盘点规则", "一物一行，版本、数量、位置、责任人必须可追溯"), ("缺口处理", "未到位资源需标注影响、责任方和计划日期"), ("交付判定", "影响联调或上线的缺口不得标记为已确认")],
    "catalog": [("目录规则", "同一业务的不同版本分别登记并保留样例"), ("打印验证", "记录纸张、方向、边距、条码与打印机"), ("发布规则", "业务确认后锁定版本并归档变更说明")],
    "countdown": [("指挥链", "总负责人、业务负责人、技术负责人和应急联系人到位"), ("升级机制", "阻塞事项按时限升级，不在备注中静默延期"), ("放行依据", "范围、环境、数据、人员和回退条件全部满足")],
    "deployment": [("变更依据", "每次执行关联变更单、制品标识和部署窗口"), ("敏感信息", "口令与密钥仅填写安全保管位置，不在文档明文记录"), ("可回退", "执行前确认备份位置、回退步骤和验证负责人")],
    "configuration": [("配置来源", "每项配置关联需求、数据字典或确认记录"), ("双人复核", "关键参数由非配置人员复核并留下证据"), ("变更留痕", "导出配置基线，后续变更另存版本")],
    "technical-review": [("文档基线", "确认接口文档、字段字典和异常码版本"), ("边界确认", "明确系统职责、网络、安全和日志边界"), ("未决事项", "形成责任人、决定日期与临时处理方案")],
    "diagram-review": [("图示约定", "参与方、生命线、同步/异步调用及数据落点清晰"), ("异常覆盖", "超时、重试、失败、回退和人工补偿均有分支"), ("一致性", "图示与接口文档、业务流程及部署方案一致")],
    "integration": [("执行原则", "每次联调有批次、有数据引用、有链路标识"), ("缺陷闭环", "失败场景关联问题编号、修复版本和复测结论"), ("完成标准", "正常、异常和边界场景均由双方确认")],
    "acceptance": [("签署前提", "关键用例通过且阻断问题关闭"), ("有条件通过", "必须写明限制条件、责任人和关闭日期"), ("不通过", "保留证据并明确重新验收窗口")],
    "test-charter": [("用例优先级", "先验证高频、关键、跨系统与高风险场景"), ("数据安全", "使用脱敏数据，截图和日志不得含真实身份信息"), ("退出标准", "阻断和严重问题关闭，其余问题有可接受计划")],
    "training": [("分组培训", "按岗位配置权限、操作路径和常见问题"), ("效果验证", "关键岗位需完成实操或理解度考核"), ("资料归档", "课件、签到、问答与补训计划统一留档")],
    "trial-observation": [("现场记录", "按科室或岗位记录真实业务运行情况"), ("隐私保护", "仅使用脱敏标识或问题编号描述异常"), ("收口判断", "按稳定性、可用性和遗留风险形成结论")],
    "cutover-command": [("单一指挥", "上线窗口内由总指挥统一决策和通知"), ("分钟级执行", "步骤按时间排序，每项均有完成标志"), ("回退触发", "触发条件、决策人、执行人和验证方式明确")],
    "formal-confirmation": [("结论分级", "通过、有条件通过、不通过三选一"), ("限制条件", "有条件通过必须登记责任与关闭期限"), ("多方确认", "技术、业务和管理责任人分别签署")],
}


def add_metric_strip(doc: Document, spec: dict) -> None:
    accent = spec["accent"]
    metrics = OVERVIEW_METRICS[spec["layout"]]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = normalize_widths([1, 1, 1, 1])
    set_table_geometry(table, widths)
    set_table_borders(table, color=WHITE, size="0")
    for index, (label, detail) in enumerate(metrics):
        cell = table.cell(0, index)
        set_cell_fill(cell, accent if spec["layout"] in {"countdown", "cutover-command"} else accent_fill(accent))
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(f"{label}\n")
        set_run_font(run, size=12, color=WHITE if spec["layout"] in {"countdown", "cutover-command"} else accent, bold=True)
        run = paragraph.add_run(detail)
        set_run_font(run, size=8.5, color=WHITE if spec["layout"] in {"countdown", "cutover-command"} else BODY)


def add_overview_panel(doc: Document, spec: dict) -> None:
    accent = spec["accent"]
    add_section_title(doc, f"一、{LAYOUT_LABELS[spec['layout']]}总览", accent)
    add_metric_strip(doc, spec)
    add_summary_table(doc, spec["summary_labels"])
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    widths = normalize_widths([1.55, 8.45])
    set_table_geometry(table, widths)
    set_table_borders(table, color=accent)
    for index, (label, detail) in enumerate(OVERVIEW_DETAILS[spec["layout"]]):
        set_cell_fill(table.cell(index, 0), accent)
        set_cell_fill(table.cell(index, 1), accent_fill(accent))
        set_cell_text(table.cell(index, 0), label, size=9, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(index, 1), detail, size=9, color=BODY)


def add_main_table(doc: Document, headers: Sequence[str], weights: Sequence[float], row_count: int,
                   accent: str = PRIMARY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = normalize_widths(weights)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_fill(table.cell(0, index), NAVY if index == 0 else accent)
        set_cell_text(table.cell(0, index), header, size=8.5, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_number in range(1, row_count + 1):
        row = table.add_row()
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_fill(cell, WHITE if index == 0 else accent_fill(accent))
            set_cell_text(cell, str(row_number) if index == 0 else "", size=8.5,
                          color=MUTED if index == 0 else BODY,
                          align=WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT)
        set_table_geometry(table, widths)


def add_diagram_box(doc: Document, accent: str = PURPLE) -> None:
    add_section_title(doc, "二、业务时序图画布", accent, page_break_before=True)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent)
    cell = table.cell(0, 0)
    set_cell_fill(cell, accent_fill(accent))
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(105)
    run = p.add_run("在此粘贴已编号的时序图，或填写附件名称与归档链接")
    set_run_font(run, size=10, color=MUTED, italic=True)


CLOSING_CONFIG = {
    "inventory": ("资源缺口与到位跟踪", ["缺口项", "业务影响", "责任方", "计划到位/证据"]),
    "catalog": ("样例验证与版本发布", ["模板/样例", "验证场景", "确认意见", "发布版本/日期"]),
    "countdown": ("阻塞事项与放行决策", ["阻塞事项", "影响窗口", "升级责任人", "处置/决策"]),
    "deployment": ("部署验证与回退记录", ["验证对象", "验证结果/证据", "回退触发条件", "执行人/时间"]),
    "configuration": ("差异复核与配置基线", ["配置差异", "影响评估", "复核结论", "基线文件/版本"]),
    "technical-review": ("未决技术事项与版本锁定", ["议题", "方案/决定", "责任人", "完成日期/版本"]),
    "diagram-review": ("评审决定与异常分支", ["评审点", "决定/异常处理", "责任人", "完成日期/状态"]),
    "integration": ("缺陷闭环与联调移交", ["问题编号", "根因/修复版本", "复测结论", "双方确认"]),
    "acceptance": ("验收限制条件与关闭计划", ["限制条件", "业务影响", "责任人", "关闭日期/复验"]),
    "test-charter": ("缺陷与测试退出判定", ["问题/缺陷", "等级与影响", "修复/复测", "退出结论"]),
    "training": ("培训效果与补训安排", ["岗位/人员", "考核或反馈", "薄弱项", "补训责任/日期"]),
    "trial-observation": ("现场观察与上线准备度", ["观察事项", "影响/频次", "跟进措施", "关闭情况"]),
    "cutover-command": ("上线风险、回退与通知", ["风险/触发事件", "应对或回退步骤", "决策/执行人", "通知对象/结果"]),
    "formal-confirmation": ("有条件事项与最终结论", ["限制/遗留事项", "业务影响", "责任方", "关闭日期/状态"]),
}


SIGNOFF_ROLES = {
    "inventory": ["实施盘点人", "客户信息部门", "项目经理"],
    "catalog": ["模板维护人", "业务代表", "项目经理"],
    "countdown": ["实施负责人", "客户项目负责人", "上线总指挥"],
    "deployment": ["部署执行人", "复核/验证人", "变更负责人"],
    "configuration": ["配置执行人", "独立复核人", "业务确认人"],
    "technical-review": ["接口负责人", "研发负责人", "客户信息部门"],
    "diagram-review": ["图示编制人", "评审主持人", "业务/技术代表"],
    "integration": ["实施方联调人", "对方联调人", "双方负责人"],
    "acceptance": ["实施方负责人", "客户技术负责人", "客户项目负责人"],
    "test-charter": ["测试负责人", "业务代表", "项目经理"],
    "training": ["培训讲师", "组织负责人", "客户代表"],
    "trial-observation": ["现场负责人", "业务代表", "项目经理"],
    "cutover-command": ["执行负责人", "业务负责人", "上线总指挥"],
    "formal-confirmation": ["技术负责人", "业务负责人", "管理负责人"],
}


def add_closing_panel(doc: Document, spec: dict, start_new_page: bool = False) -> None:
    accent = spec["accent"]
    title, headers = CLOSING_CONFIG[spec["layout"]]
    closing_number = "四" if spec.get("diagram_box") else "三"
    gate_number = "五" if spec.get("diagram_box") else "四"
    add_section_title(doc, f"{closing_number}、{title}", accent, page_break_before=start_new_page)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    widths = normalize_widths([1.8, 3.3, 1.5, 3.4])
    set_table_geometry(table, widths)
    set_table_borders(table, color=accent)
    for index, header in enumerate(headers):
        set_cell_fill(table.cell(0, index), accent)
        set_cell_text(table.cell(0, index), header, size=8.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_index in range(1, 5):
        for col_index in range(4):
            set_cell_fill(table.cell(row_index, col_index), WHITE if col_index == 0 else accent_fill(accent))
            set_cell_text(table.cell(row_index, col_index), "", size=8.7)

    add_section_title(doc, f"{gate_number}、交付门禁", accent)
    gates = doc.add_table(rows=1, cols=3)
    gates.style = "Table Grid"
    set_table_geometry(gates, normalize_widths([1, 1, 1]))
    set_table_borders(gates, color=accent)
    for index, (label, detail) in enumerate(spec["checks"]):
        cell = gates.cell(0, index)
        set_cell_fill(cell, accent_fill(accent))
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(f"{label}\n")
        set_run_font(run, size=9.2, color=accent, bold=True)
        run = paragraph.add_run(f"{detail}\n☐ 通过  ☐ 待整改  ☐ 不适用")
        set_run_font(run, size=8.6, color=BODY)


def add_checklist(doc: Document, checks: Iterable[tuple[str, str]], heading_number: str) -> None:
    add_section_title(doc, f"{heading_number}、质量核对")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = normalize_widths([0.6, 1.25, 4.90, 1.35])
    set_table_geometry(table, widths)
    set_table_borders(table)
    for index, header in enumerate(["序号", "检查维度", "核对要求", "结论"]):
        set_cell_fill(table.cell(0, index), NAVY)
        set_cell_text(table.cell(0, index), header, size=8.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_number, (label, detail) in enumerate(checks, start=1):
        row = table.add_row()
        prevent_row_split(row)
        values = [str(row_number), label, detail, "☐ 通过  ☐ 不通过  ☐ 不适用"]
        for index, value in enumerate(values):
            set_cell_fill(row.cells[index], WHITE if index < 3 else PALE_YELLOW)
            set_cell_text(row.cells[index], value, size=8.8, color=BODY,
                          align=WD_ALIGN_PARAGRAPH.CENTER if index in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)


def add_conclusion(doc: Document, heading_number: str) -> None:
    add_section_title(doc, f"{heading_number}、结论与遗留事项")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    widths = normalize_widths([1.10, 8.90])
    set_table_geometry(table, widths)
    set_table_borders(table)
    labels = ["总体结论", "遗留事项/风险", "附件与证据"]
    placeholders = ["☐ 通过  ☐ 有条件通过  ☐ 不通过", "", ""]
    for index, label in enumerate(labels):
        set_cell_fill(table.cell(index, 0), PALE_BLUE)
        set_cell_fill(table.cell(index, 1), PALE_YELLOW)
        set_cell_text(table.cell(index, 0), label, size=9, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(index, 1), placeholders[index], size=9, color=BODY)


def add_signoff(doc: Document, heading_number: str, spec: dict) -> None:
    accent = spec["accent"]
    add_section_title(doc, f"{heading_number}、责任确认", accent, page_break_before=True)
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    widths = normalize_widths([1.25, 1.80, 5.45, 1.50])
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["确认角色", "姓名/签字", "确认意见", "日期"]
    for index, header in enumerate(headers):
        set_cell_fill(table.cell(0, index), accent)
        set_cell_text(table.cell(0, index), header, size=8.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    roles = SIGNOFF_ROLES[spec["layout"]]
    for row_index, role in enumerate(roles, start=1):
        for col_index in range(4):
            set_cell_fill(table.cell(row_index, col_index), accent_fill(accent) if col_index else PALE_BLUE)
            set_cell_text(table.cell(row_index, col_index), role if col_index == 0 else "", size=8.8,
                          color=NAVY if col_index == 0 else BODY,
                          bold=col_index == 0,
                          align=WD_ALIGN_PARAGRAPH.CENTER if col_index != 2 else WD_ALIGN_PARAGRAPH.LEFT)


def add_revision_log(doc: Document, spec: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"版本记录：V2.1 · {LAYOUT_LABELS[spec['layout']]}差异化改版 · 2026-07-28")
    set_run_font(run, size=8, color=MUTED)


def build_document(spec: dict) -> Path:
    doc = Document()
    configure_document(doc, spec)
    add_title_block(doc, spec)
    add_callout(doc, spec["instruction"], spec["accent"])
    add_overview_panel(doc, spec)

    if spec.get("diagram_box"):
        add_diagram_box(doc, spec["accent"])
        main_heading = "三、评审记录"
    else:
        main_heading = {
            "inventory": "二、资源盘点明细",
            "catalog": "二、模板目录与版本",
            "countdown": "二、倒排任务计划",
            "deployment": "二、部署执行日志",
            "configuration": "二、配置项与验证",
            "technical-review": "二、接口文档评审矩阵",
            "integration": "二、联调执行记录",
            "acceptance": "二、接口验收矩阵",
            "test-charter": "二、端到端测试用例",
            "training": "二、参训与效果记录",
            "trial-observation": "二、现场试行记录",
            "cutover-command": "二、上线执行清单",
            "formal-confirmation": "二、上线确认矩阵",
        }[spec["layout"]]

    add_section_title(doc, main_heading, spec["accent"], page_break_before=True)
    add_main_table(doc, spec["columns"], spec["weights"], spec["rows"], spec["accent"])
    add_closing_panel(doc, spec, start_new_page=not spec.get("diagram_box"))
    add_signoff(doc, "六" if spec.get("diagram_box") else "五", spec)
    add_revision_log(doc, spec)

    output = OUTPUT_DIR / f"{WORD_FILENAMES.get(spec['slug'], spec['slug'])}.docx"
    doc.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_document(spec) for spec in SPECS]
    print(f"generated {len(outputs)} Word templates")
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
